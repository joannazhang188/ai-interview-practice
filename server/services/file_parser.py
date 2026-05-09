import os
import re
import fitz
from docx import Document

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        return f"[PDF解析失败: {str(e)}]"
    return text.strip()

def extract_text_from_docx(file_path: str) -> str:
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        return f"[Word解析失败: {str(e)}]"
    return text.strip()

def extract_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    elif ext in [".png", ".jpg", ".jpeg"]:
        return "[图片文件暂不支持解析，请上传文字版简历]"
    else:
        return "[不支持的文件格式]"

COMMON_SKILLS = [
    "Python", "Java", "JavaScript", "TypeScript", "Go", "Rust", "C++", "C#",
    "React", "Vue", "Angular", "Node.js", "Django", "Flask", "Spring",
    "SQL", "MySQL", "PostgreSQL", "MongoDB", "Redis",
    "AWS", "Azure", "GCP", "Docker", "Kubernetes", "Linux",
    "Git", "Agile", "Scrum",
    "产品设计", "需求分析", "项目管理", "数据分析", "用户研究",
    "Axure", "Figma", "Sketch", "SQL", "Excel", "PPT",
    "机器学习", "深度学习", "NLP", "计算机视觉",
    "沟通能力", "团队协作", "领导力", "问题解决", "逻辑思维",
    "Java Spring Boot", "MySQL", "Redis", "ES", "Kafka",
    "用户增长", "活动策划", "内容运营", "渠道运营", "数据可视化",
    "B端产品", "C端产品", "AI产品", "策略产品", "数据产品",
]

def extract_skills(text: str) -> list:
    found = []
    text_upper = text.upper()
    for skill in COMMON_SKILLS:
        if skill.upper() in text_upper:
            found.append(skill)
    return list(set(found))[:20]

def extract_education(text: str) -> list:
    edu_keywords = ["本科", "硕士", "博士", "学士", "研究生", "大学", "学院", "毕业"]
    lines = text.split("\n")
    education = []
    for line in lines:
        for kw in edu_keywords:
            if kw in line and len(line) < 100:
                education.append(line.strip())
                break
    return education[:5]

def extract_experience(text: str) -> list:
    exp_keywords = ["负责", "主导", "参与", "担任", "项目", "产品", "设计", "开发", "运营", "管理"]
    lines = text.split("\n")
    experience = []
    for line in lines:
        if any(kw in line for kw in exp_keywords) and 5 < len(line) < 200:
            experience.append(line.strip())
    return experience[:10]

def parse_resume(text: str) -> dict:
    return {
        "skills": extract_skills(text),
        "education": extract_education(text),
        "experience": extract_experience(text),
        "summary": text[:300] + "..." if len(text) > 300 else text
    }

def parse_portfolio(text: str) -> dict:
    return {
        "summary": text[:400] + "..." if len(text) > 400 else text,
        "skills": extract_skills(text),
        "highlights": extract_experience(text),
        "suggested_roles": []
    }
